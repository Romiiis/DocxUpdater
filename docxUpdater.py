# Imports
from docx import Document
import re
from tkinter import filedialog, messagebox

# Prefixes for the print statements
prefix_info = "[INFO] "
prefix_error = "[ERROR] "

# Dialog box to select the docx file
filename = ""

# If the filename is empty, show an error message and exit the program or if the file is not a docx file (try 2 times)
for i in range(2):

    # Open the file dialog box to select the docx file
    filename = filedialog.askopenfilename()

    # If the filename is empty, show an error message and exit the program
    if filename == '':
        messagebox.showerror('Error', 'No file selected \nPlease select a .docx file')
    # If the file is not a docx file, show an error message and ask to select a docx file
    elif not filename.endswith('.docx'):
        messagebox.showerror('Error', 'Please select a docx file')
    # If the file is a docx file, break the loop
    else:
        break

# If the filename is empty, show an error message and exit the program
if filename == "":
    messagebox.showerror('Error', 'Operation cancelled')
    exit()

try:
    # Create a Document object from the docx file
    doc = Document(filename)

    # Find the first year in the docx file and store it in the variable year
    year = None
    for para in doc.paragraphs:
        if re.search(r'\b\d{4}\b', para.text):
            year = re.search(r'\b\d{4}\b', para.text).group()
            break

    # Find the year in the table and replace it with the next year
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if re.search(r'\b\d{4}\b', cell.text):
                    for match in re.finditer(r'\b\d{4}\b', cell.text):
                        if match.group() == year:
                            cell.text = cell.text.replace(year, str(int(year) + 1))

    # Find the year in the paragraphs and replace it with the next year
    for i in range(1, -2, -1):
        for para in doc.paragraphs:
            if re.search(r'\b\d{4}\b', para.text):
                for match in re.finditer(r'\b\d{4}\b', para.text):
                    if match.group() == str(int(year) + i):
                        para.text = para.text.replace(str(int(year) + i), str(int(year) + i + 1))

    # Find the year in the filename and replace it with the next year
    filename = filename.replace(year, str(int(year) + 1))

    # Save the updated docx file
    doc.save(filename.split('/')[-1])

    # Dialog box to show the file is updated and saved
    messagebox.showinfo('Info', 'File updated and saved as ' + filename.split('/')[-1])

except Exception as e:
    # If any error occurs, show an error message and exit the program
    print("Error: " + str(e))
    messagebox.showerror('Error', 'An error occurred \nPlease try again or contact the developer')
    input("Press ENTER key to exit")
    exit()
